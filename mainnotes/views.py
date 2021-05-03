from django.shortcuts import render, redirect
from .models import Notes, Image, HiddenNotePassword
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse

from docx import *
from docx.shared import Inches
from io import BytesIO

from reportlab.pdfgen import canvas
# Create your views here.


def mainnotes(request):
    if request.user.is_authenticated:
        allNotes = Notes.objects.filter(owner = request.user, hidden_note = False, trashed_note = False)
        return render(request, 'mainnotes/mainnotes.html', context = {'all_notes': allNotes})
    else:
        return redirect('login')
    
@login_required(redirect_field_name='login')
def addNote(request):
    if request.method == "POST":
        title = request.POST['title']
        description = request.POST['description']
        owner = request.user
        
        note = Notes(owner=owner, title=title, description=description)
        note.save()
        
        images = request.FILES.getlist('images')
        print(images)
        print(request.FILES)
        for image in images:
            print("Entered")
            note_image = Image.objects.create(note=note,image=image)
            note_image.save()
        
        images_for_html = Image.objects.filter(note=note)
        return redirect('/editnote/'+str(note.pk))
    
    else:
        return render(request, "mainnotes/addnote.html")
   
@login_required(redirect_field_name='login') 
def editnote(request, pk):
    if request.method == 'POST':
        title = request.POST['title']
        print("Printing title:", title)
        description = request.POST['description']
            
        note = Notes.objects.get(pk=pk)
        note.title = title
        note.description = description
        note.save()
        
        return redirect('/editnote/'+str(note.pk))
    
    else:
        note = Notes.objects.get(pk = pk)
        images = Image.objects.filter(note = note)
        return render(request, 'mainnotes/editnote.html', context={'note':note, 'images':images})



@login_required(redirect_field_name='login')
def empty_trash(request):
    notes = Notes.objects.filter(owner=request.user, trashed_note=True)
    notes.delete()
    messages.info(request,'Trash emptied successfully!')
    return redirect('/trashednotes')
    
        
@login_required(redirect_field_name='login')
def trashnote(request, pk):
    note = Notes.objects.get(pk = pk)
    if note.trashed_note == False:
        note.trashed_note = True
        note.save()
        messages.info(request,"Note Trashed Successfully")
        return redirect('/')
    
    else:
        note.trashed_note = False
        note.save()
        messages.info(request,"Note Restored Successfully")
        return redirect('/trashednotes')
    
@login_required(redirect_field_name='login')
def trashed_notes(request):
    allNotes = Notes.objects.filter(owner = request.user, trashed_note = True)
    return render(request, 'mainnotes/trashednotes.html', context = {'all_notes': allNotes})



@login_required(redirect_field_name='login')
def hidenote(request, pk):
    note = Notes.objects.get(pk = pk)
    if note.hidden_note == False:
        note.hidden_note = True
        note.save()
        messages.info(request,"Note Hided Successfully")
        return redirect('/editnote/'+str(note.pk))
    
    else:
        note.hidden_note = False
        note.save()
        messages.info(request,"Note Unhided Successfully")
        return redirect('/editnote/'+str(note.pk))
    

@login_required(redirect_field_name='login')
def hidden_notes(request):
    if HiddenNotePassword.objects.filter(user=request.user).exists():
        if request.method == 'POST':
            passcode = request.POST['passcode']
            if HiddenNotePassword.objects.get(user=request.user).hidden_note_password == passcode:
                allNotes = Notes.objects.filter(owner = request.user, hidden_note = True, trashed_note = False)
                return render(request, 'mainnotes/hiddennotes.html', context = {'all_notes': allNotes})
            else:
                messages.info(request, 'Invalid passcode for hidden notes')
                return redirect('hiddennotes')
            
        else:
            return render(request, 'mainnotes/enterhiddenpass.html')
        
    else:
        if request.method == 'POST':
            passcode = request.POST['passcode']
            passcode2 = request.POST['passcode2']
            if passcode == passcode2:
                user = request.user
                hiddenpasscode = HiddenNotePassword(user=user, hidden_note_password = passcode)
                hiddenpasscode.save()
                messages.info(request, 'Created new hidden passcode. You can view your hidden notes using it.')
            else:
                messages.info(request, 'The passcodes do not match!')
            
            return redirect('hiddennotes')
            
        else:
            return render(request, 'mainnotes/createhiddenpass.html')


@login_required(redirect_field_name='login')
def addHiddenNote(request):
    if request.method == "POST":
        title = request.POST['title']
        description = request.POST['description']
        owner = request.user
        
        note = Notes(owner=owner, title=title, description=description, hidden_note=True)
        note.save()
        
        images = request.FILES.getlist('images')
        print(images)
        print(request.FILES)
        for image in images:
            print("Entered")
            note_image = Image.objects.create(note=note,image=image)
            note_image.save()
        
        images_for_html = Image.objects.filter(note=note)
        return redirect('/editnote/'+str(note.pk))
    
    else:
        return render(request, "mainnotes/addhiddennote.html")
    
    
def export_to_docx(request, pk):
    note = Notes.objects.get(pk=pk)
    
    if note.owner == request.user:
        document = Document()
        docx_title = note.title + ".docx"
        
        document.add_heading(note.title, level=1)
        document.add_paragraph()
        document.add_paragraph(note.description)
        
        # Prepare document for download        
        # -----------------------------
        f = BytesIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_title
        response['Content-Length'] = length
        
        return response
    
    
    else:
        return HttpResponse('You are not authorized for this')
    
    
def export_to_pdf(request, pk):
    note = Notes.objects.get(pk=pk)
    
    if note.owner == request.user:
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'inline; filename=' + note.title + '.pdf'

        buffer = BytesIO()
        p = canvas.Canvas(buffer)

        # Start writing the PDF here
        p.setFont("Helvetica", 26)
        p.drawString(50, 800, note.title)
        p.line(50,790,500,790)
        p.setFont("Helvetica", 14)
        p.drawString(50, 760, note.description)
        # End writing

        p.showPage()
        p.save()

        pdf = buffer.getvalue()
        buffer.close()
        response.write(pdf)

        return response
    
    else:
        return HttpResponse('You are not authorized for this action!')